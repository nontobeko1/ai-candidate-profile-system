import os
from fpdf import FPDF
from docx import Document


def create_sample_pdf(filename):
    """Create a sample PDF resume for testing"""
    pdf = FPDF()
    pdf.add_page()

    # Title
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, "John Doe - Software Engineer", ln=True, align='C')
    pdf.ln(10)

    # Contact Info
    pdf.set_font("Arial", '', 12)
    pdf.cell(200, 10, "Email: john.doe@email.com | Phone: +1-555-0123 | Location: New York, USA", ln=True)
    pdf.ln(10)

    # Summary
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, "Professional Summary", ln=True)
    pdf.set_font("Arial", '', 12)
    pdf.multi_cell(0, 10, "Experienced Software Engineer with 5+ years in full-stack web development. "
                          "Proficient in Python, JavaScript, React, Node.js, and AWS. "
                          "Strong background in building scalable applications and leading development teams.")
    pdf.ln(10)

    # Skills
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, "Technical Skills", ln=True)
    pdf.set_font("Arial", '', 12)
    skills = [
        "Programming: Python, JavaScript, Java, C++, SQL",
        "Frameworks: React, Node.js, Django, Flask, Spring Boot",
        "Tools: Git, Docker, AWS, Azure, Kubernetes, Jenkins",
        "Databases: MySQL, MongoDB, PostgreSQL, Redis"
    ]
    for skill in skills:
        pdf.cell(200, 10, skill, ln=True)
    pdf.ln(10)

    # Experience
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, "Work Experience", ln=True)

    pdf.set_font("Arial", 'B', 12)
    pdf.cell(200, 10, "Senior Software Engineer - Tech Solutions Inc.", ln=True)
    pdf.set_font("Arial", 'I', 12)
    pdf.cell(200, 10, "2020 - Present | New York, USA", ln=True)
    pdf.set_font("Arial", '', 12)
    pdf.multi_cell(0, 10, "Led development of microservices architecture. Improved system performance by 40%. "
                          "Technologies: Python, React, AWS, Docker")
    pdf.ln(5)

    pdf.set_font("Arial", 'B', 12)
    pdf.cell(200, 10, "Software Developer - Innovation Labs", ln=True)
    pdf.set_font("Arial", 'I', 12)
    pdf.cell(200, 10, "2018 - 2020 | Boston, USA", ln=True)
    pdf.set_font("Arial", '', 12)
    pdf.multi_cell(0, 10, "Developed web applications using JavaScript and Python. "
                          "Collaborated with cross-functional teams. Technologies: Django, Vue.js, MySQL")
    pdf.ln(10)

    # Education
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, "Education", ln=True)

    pdf.set_font("Arial", 'B', 12)
    pdf.cell(200, 10, "Master of Science in Computer Science", ln=True)
    pdf.set_font("Arial", 'I', 12)
    pdf.cell(200, 10, "Stanford University | 2016 - 2018 | GPA: 3.8/4.0", ln=True)
    pdf.ln(5)

    pdf.set_font("Arial", 'B', 12)
    pdf.cell(200, 10, "Bachelor of Science in Software Engineering", ln=True)
    pdf.set_font("Arial", 'I', 12)
    pdf.cell(200, 10, "MIT | 2012 - 2016 | GPA: 3.6/4.0", ln=True)

    pdf.output(filename)
    print(f"✅ Created sample PDF: {filename}")


def create_sample_docx(filename):
    """Create a sample DOCX resume for testing"""
    doc = Document()

    # Title
    title = doc.add_heading('John Doe - Software Engineer', 0)
    title.alignment = 1  # Center alignment

    # Contact Info
    doc.add_paragraph('Email: john.doe@email.com | Phone: +1-555-0123 | Location: New York, USA')
    doc.add_paragraph()

    # Summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph('Experienced Software Engineer with 5+ years in full-stack web development. '
                      'Proficient in Python, JavaScript, React, Node.js, and AWS. '
                      'Strong background in building scalable applications and leading development teams.')
    doc.add_paragraph()

    # Skills
    doc.add_heading('Technical Skills', level=1)
    skills = [
        "Programming: Python, JavaScript, Java, C++, SQL, HTML/CSS",
        "Frameworks: React, Angular, Node.js, Django, Flask, Spring Boot",
        "Tools: Git, Docker, AWS, Azure, Kubernetes, Jenkins, Jira",
        "Databases: MySQL, MongoDB, PostgreSQL, Redis, Elasticsearch",
        "Soft Skills: Leadership, Communication, Problem Solving, Teamwork, Project Management"
    ]
    for skill in skills:
        doc.add_paragraph(skill, style='List Bullet')
    doc.add_paragraph()

    # Experience
    doc.add_heading('Work Experience', level=1)

    # Job 1
    doc.add_heading('Senior Software Engineer - Tech Solutions Inc.', level=2)
    doc.add_paragraph('2020 - Present | New York, USA')
    doc.add_paragraph('Led development of microservices architecture serving 1M+ users. '
                      'Improved system performance by 40% through optimization. '
                      'Technologies used: Python, React, AWS, Docker, Kubernetes')
    doc.add_paragraph()

    # Job 2
    doc.add_heading('Software Developer - Innovation Labs', level=2)
    doc.add_paragraph('2018 - 2020 | Boston, USA')
    doc.add_paragraph('Developed and maintained web applications using JavaScript and Python. '
                      'Collaborated with cross-functional teams in agile environment. '
                      'Technologies: Django, Vue.js, MySQL, Redis')
    doc.add_paragraph()

    # Education
    doc.add_heading('Education', level=1)

    doc.add_heading('Master of Science in Computer Science', level=2)
    doc.add_paragraph('Stanford University | 2016 - 2018 | GPA: 3.8/4.0')
    doc.add_paragraph('Thesis: Machine Learning Applications in Web Security')
    doc.add_paragraph()

    doc.add_heading('Bachelor of Science in Software Engineering', level=2)
    doc.add_paragraph('MIT | 2012 - 2016 | GPA: 3.6/4.0')
    doc.add_paragraph('Minor in Business Administration')

    doc.save(filename)
    print(f"✅ Created sample DOCX: {filename}")


def create_sample_transcript_pdf(filename):
    """Create a sample academic transcript PDF"""
    pdf = FPDF()
    pdf.add_page()

    # Title
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, "UNIVERSITY OF TECHNOLOGY", ln=True, align='C')
    pdf.cell(200, 10, "ACADEMIC TRANSCRIPT", ln=True, align='C')
    pdf.ln(10)

    # Student Info
    pdf.set_font("Arial", '', 12)
    pdf.cell(200, 10, "Student: John Doe", ln=True)
    pdf.cell(200, 10, "Student ID: 12345678", ln=True)
    pdf.cell(200, 10, "Program: Bachelor of Computer Science", ln=True)
    pdf.cell(200, 10, "Graduation Date: June 2018", ln=True)
    pdf.ln(10)

    # Grades Table Header
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(60, 10, "Course Code", 1)
    pdf.cell(80, 10, "Course Name", 1)
    pdf.cell(30, 10, "Credits", 1)
    pdf.cell(30, 10, "Grade", 1, ln=True)

    # Grades
    pdf.set_font("Arial", '', 12)
    grades = [
        ("CS101", "Introduction to Programming", "4", "85"),
        ("CS102", "Data Structures", "4", "88"),
        ("CS201", "Algorithms", "4", "92"),
        ("CS301", "Database Systems", "4", "87"),
        ("CS302", "Web Development", "4", "90"),
        ("CS401", "Machine Learning", "4", "94"),
        ("CS402", "Software Engineering", "4", "89"),
        ("MATH101", "Calculus I", "3", "83"),
        ("MATH102", "Linear Algebra", "3", "86"),
        ("STAT201", "Probability & Statistics", "3", "88")
    ]

    for course in grades:
        pdf.cell(60, 10, course[0], 1)
        pdf.cell(80, 10, course[1], 1)
        pdf.cell(30, 10, course[2], 1)
        pdf.cell(30, 10, course[3], 1, ln=True)

    pdf.ln(10)

    # Summary
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(200, 10, "Cumulative GPA: 3.7 / 4.0", ln=True)
    pdf.cell(200, 10, "Classification: First Class Honors", ln=True)
    pdf.cell(200, 10, "Total Credits: 42", ln=True)

    pdf.output(filename)
    print(f"✅ Created sample transcript PDF: {filename}")


def create_sample_transcript_docx(filename):
    """Create a sample academic transcript DOCX"""
    doc = Document()

    # Title
    title = doc.add_heading('UNIVERSITY OF TECHNOLOGY', 0)
    title.alignment = 1
    doc.add_heading('ACADEMIC TRANSCRIPT', level=1).alignment = 1

    # Student Info
    doc.add_paragraph('Student: John Doe')
    doc.add_paragraph('Student ID: 12345678')
    doc.add_paragraph('Program: Bachelor of Computer Science')
    doc.add_paragraph('Graduation Date: June 2018')
    doc.add_paragraph()

    # Grades Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Course Code'
    hdr_cells[1].text = 'Course Name'
    hdr_cells[2].text = 'Credits'
    hdr_cells[3].text = 'Grade'

    # Grades data
    grades = [
        ("CS101", "Introduction to Programming", "4", "85"),
        ("CS102", "Data Structures", "4", "88"),
        ("CS201", "Algorithms", "4", "92"),
        ("CS301", "Database Systems", "4", "87"),
        ("CS302", "Web Development", "4", "90"),
        ("CS401", "Machine Learning", "4", "94"),
        ("CS402", "Software Engineering", "4", "89"),
        ("MATH101", "Calculus I", "3", "83"),
        ("ENG101", "Technical Writing", "3", "85")
    ]

    for course in grades:
        row_cells = table.add_row().cells
        row_cells[0].text = course[0]
        row_cells[1].text = course[1]
        row_cells[2].text = course[2]
        row_cells[3].text = course[3]

    doc.add_paragraph()

    # Summary
    doc.add_paragraph('Cumulative GPA: 3.7 / 4.0')
    doc.add_paragraph('Classification: First Class Honors')
    doc.add_paragraph('Total Credits: 39')

    doc.save(filename)
    print(f"✅ Created sample transcript DOCX: {filename}")


if __name__ == "__main__":
    print("📄 Creating sample test documents...")
    print("=" * 50)

    # Create sample resumes
    create_sample_pdf("sample_resume.pdf")
    create_sample_docx("sample_resume.docx")

    # Create sample transcripts
    create_sample_transcript_pdf("sample_transcript.pdf")
    create_sample_transcript_docx("sample_transcript.docx")

    print("=" * 50)
    print("🎉 All sample documents created successfully!")
    print("\nNow you can test the document processor:")
    print("python document_processor.py")